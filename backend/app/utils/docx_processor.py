"""
Utility for processing DOCX files and extracting their content.
"""
import io
import os
import logging
import re
from typing import Dict, Any, Optional, List, Tuple
import docx
from docx import Document
from docx.table import Table, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from ..core.logging import service_logger as logger


class DocxProcessor:
    """Class for processing DOCX files and extracting their content."""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from a DOCX file including tables and other elements.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            String containing all text from the document
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = DocxProcessor._get_cell_text(cell)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            result = "\n".join(full_text)
            logger.debug(f"Extracted {len(result)} characters from DOCX file")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
            return ""
    
    @staticmethod
    def _get_cell_text(cell) -> str:
        """
        Extract text from a table cell, including nested tables.
        
        Args:
            cell: DOCX table cell object
            
        Returns:
            String containing all text from the cell
        """
        text_parts = []
        
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Handle nested tables
        for table in cell.tables:
            for row in table.rows:
                row_text = []
                for nested_cell in row.cells:
                    nested_text = DocxProcessor._get_cell_text(nested_cell)
                    if nested_text:
                        row_text.append(nested_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return " ".join(text_parts)
    
    @staticmethod
    def extract_qa_pairs(text: str) -> Dict[str, Dict[str, Any]]:
        """
        Extract question-answer pairs from document text using multiple patterns.
        
        Args:
            text: The document text to process
            
        Returns:
            Dictionary of question-answer pairs
        """
        try:
            # Try section-based extraction first (for academic papers)
            section_qa_pairs = DocxProcessor._extract_qa_from_sections(text)
            if section_qa_pairs:
                logger.info(f"Extracted {len(section_qa_pairs)} Q&A pairs using section-based method")
                return section_qa_pairs
            
            # Try pattern-based extraction with multiple regex patterns
            patterns_qa_pairs = DocxProcessor._extract_qa_with_patterns(text)
            if patterns_qa_pairs:
                logger.info(f"Extracted {len(patterns_qa_pairs)} Q&A pairs using pattern-based method")
                return patterns_qa_pairs
            
            # Try table-based extraction
            table_qa_pairs = DocxProcessor._extract_qa_from_tables(text)
            if table_qa_pairs:
                logger.info(f"Extracted {len(table_qa_pairs)} Q&A pairs using table-based method")
                return table_qa_pairs
            
            logger.warning("No Q&A pairs found using any extraction method")
            return {}
            
        except Exception as e:
            logger.error(f"Error extracting Q&A pairs: {e}")
            return {}
    
    @staticmethod
    def _extract_qa_from_sections(text: str) -> Dict[str, Dict[str, Any]]:
        """Extract Q&A pairs from section-based document structure."""
        qa_pairs = {}
        
        # Look for section headers and content
        section_pattern = r'^(?:#{1,6}\s*)?(?:Section\s*\d+|Question\s*\d+|Q\d+|Problem\s*\d+)[:\s]*(.+?)(?=^(?:#{1,6}\s*)?(?:Section\s*\d+|Question\s*\d+|Q\d+|Problem\s*\d+)|$)'
        
        sections = re.findall(section_pattern, text, re.MULTILINE | re.DOTALL | re.IGNORECASE)
        
        for i, section in enumerate(sections):
            section_text = section.strip()
            if len(section_text) > 50:  # Ensure substantial content
                # Try to split into question and answer
                lines = section_text.split('\n')
                if len(lines) >= 2:
                    question = lines[0].strip()
                    answer = '\n'.join(lines[1:]).strip()
                    
                    if question and answer:
                        qa_pairs[f"section_{i+1}"] = {
                            "question": question,
                            "answer": answer
                        }
        
        return qa_pairs
    
    @staticmethod
    def _extract_qa_with_patterns(text: str) -> Dict[str, Dict[str, Any]]:
        """Extract Q&A pairs using regex patterns."""
        qa_pairs = {}
        
        # Multiple patterns to try
        patterns = [
            # Pattern 1: Q: ... A: ...
            r'(?:^|\n)(?:Q|Question)(?:\s*\d+)?[\s:.]+(.+?)(?=\n(?:A|Answer)[\s:.]+).*?\n(?:A|Answer)(?:\s*\d+)?[\s:.]+(.+?)(?=\n(?:Q|Question)|\n\n|$)',
            
            # Pattern 2: numbered questions
            r'(?:^|\n)(\d+[\.)]\s*.+?)(?=\n(?:Answer|A[\s:.]|Solution))(.*?)(?=\n\d+[\.)]|\n\n|$)',
            
            # Pattern 3: Question/Answer headers
            r'(?:^|\n)(?:Question|Q)[\s:]*(.+?)(?=\n(?:Answer|A[\s:])).*?\n(?:Answer|A)[\s:]*(.+?)(?=\n(?:Question|Q)|\n\n|$)',
        ]
        
        for pattern_idx, pattern in enumerate(patterns):
            matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL | re.IGNORECASE)
            
            if matches:
                for i, match in enumerate(matches):
                    if len(match) == 2:
                        question, answer = match
                        question = question.strip()
                        answer = answer.strip()
                        
                        if question and answer and len(question) > 5 and len(answer) > 5:
                            qa_pairs[f"pattern_{pattern_idx+1}_q{i+1}"] = {
                                "question": question,
                                "answer": answer
                            }
                
                if qa_pairs:
                    break  # Use the first pattern that works
        
        return qa_pairs
    
    @staticmethod
    def _extract_qa_from_tables(text: str) -> Dict[str, Dict[str, Any]]:
        """Extract Q&A pairs from table structures."""
        qa_pairs = {}
        
        # Look for table-like structures with | separators
        table_pattern = r'(.+?)\s*\|\s*(.+?)(?=\n|$)'
        matches = re.findall(table_pattern, text, re.MULTILINE)
        
        question_indicators = ['question', 'q', 'problem', 'query']
        answer_indicators = ['answer', 'a', 'solution', 'response']
        
        for i, (left, right) in enumerate(matches):
            left = left.strip().lower()
            right = right.strip()
            
            # Check if left side indicates a question
            if any(indicator in left for indicator in question_indicators):
                # The right side is likely the question
                question = right
                # Look for the corresponding answer in next rows
                if i + 1 < len(matches):
                    next_left, next_right = matches[i + 1]
                    next_left = next_left.strip().lower()
                    if any(indicator in next_left for indicator in answer_indicators):
                        answer = next_right.strip()
                        
                        if question and answer and len(question) > 5 and len(answer) > 5:
                            qa_pairs[f"table_q{len(qa_pairs)+1}"] = {
                                "question": question,
                                "answer": answer
                            }
        
        return qa_pairs
    
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes) -> str:
        """
        Extract text from DOCX file bytes.
        
        Args:
            file_bytes: Bytes content of the DOCX file
            
        Returns:
            String containing all text from the document
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = DocxProcessor._get_cell_text(cell)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            result = "\n".join(full_text)
            logger.debug(f"Extracted {len(result)} characters from DOCX bytes")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {e}")
            return ""
