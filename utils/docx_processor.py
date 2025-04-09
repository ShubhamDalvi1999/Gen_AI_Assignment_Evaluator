"""
Utility for processing DOCX files and extracting their content.
"""
import io
import os
import logging
import re
from typing import Dict, Any, Optional, List, Tuple
import docx
from docx import Document
from docx.table import Table, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

logger = logging.getLogger(__name__)

class DocxProcessor:
    """Class for processing DOCX files and extracting their content."""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from a DOCX file including tables and other elements.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            String containing all text from the document
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = DocxProcessor._get_cell_text(cell)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Check if we found any text
            if not full_text:
                logger.warning(f"No text content found in {file_path}")
                return ""
                
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def _get_cell_text(cell: _Cell) -> str:
        """Extract text from a table cell, handling nested elements."""
        text = []
        for paragraph in cell.paragraphs:
            if paragraph.text:
                text.append(paragraph.text.strip())
        return "\n".join(text)
    
    @staticmethod
    def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
        """
        Extract text from DOCX file bytes including tables and other elements.
        
        Args:
            docx_bytes: Bytes of the DOCX file
            
        Returns:
            String containing all text from the document
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = DocxProcessor._get_cell_text(cell)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {e}")
            return ""
    
    # 3. Extracts Q&A pairs from text using multiple methods.
    @staticmethod
    def extract_qa_pairs(text: str) -> Dict[str, Dict[str, Any]]:
        """
        Extract question-answer pairs from text using multiple methods.
        Returns a dictionary mapping question IDs to question-answer dictionaries.
        """
        if not text:
            logging.warning("Empty text provided for Q&A extraction")
            return {}
        
        # Try section-based extraction first (for academic papers and structured documents)
        section_qa_pairs = DocxProcessor._extract_qa_from_sections(text)
        if section_qa_pairs:
            return section_qa_pairs
        
        # Fall back to pattern-based extraction
        qa_pairs = {}
        qa_count = 0
        
        # Pattern 1: Explicit Q: and A: markers
        pattern1 = r'(?:^|\n)(?:Q|Question)(?:\s*\d+)?[\s:.]+(.+?)(?=\n(?:A|Answer)[\s:.]+)(?:\n(?:A|Answer)[\s:.]+)(.+?)(?=\n(?:Q|Question)(?:\s*\d+)?[\s:.]+|\Z)'
        matches1 = re.finditer(pattern1, text, re.DOTALL | re.IGNORECASE)
        for match in matches1:
            qa_count += 1
            qa_pairs[f"qa_{qa_count}"] = {
                'question': match.group(1).strip(),
                'answer': match.group(2).strip()
            }
        
        # If explicit markers found, return those results
        if qa_pairs:
            return qa_pairs
        
        # Pattern 2: Numbered questions followed by answers
        # Extended to catch numbered questions with or without question marks
        pattern2 = r'(?:^|\n)(?:\d+[\.)]\s*|\([a-zA-Z0-9]+\)\s*|[a-zA-Z][\.)]\s*)([^?]*?(?:\?|\.|\n))(?!\d+[\.)]\s*|\([a-zA-Z0-9]+\)\s*|[a-zA-Z][\.)]\s*)(.+?)(?=(?:\n(?:\d+[\.)]\s*|\([a-zA-Z0-9]+\)\s*|[a-zA-Z][\.)]\s*))|\Z)'
        matches2 = re.finditer(pattern2, text, re.DOTALL)
        for match in matches2:
            question = match.group(1).strip()
            answer = match.group(2).strip()
            
            # Skip if the question is too short (likely a section header)
            if len(question) < 5:
                continue
            
            qa_count += 1
            qa_pairs[f"qa_{qa_count}"] = {
                'question': question,
                'answer': answer
            }
        
        # Pattern 3: Questions ending with ? followed by answers
        pattern3 = r'(?:^|\n)([^?]+\?)\s*(.+?)(?=\n[^?]+\?|\Z)'
        matches3 = re.finditer(pattern3, text, re.DOTALL)
        for match in matches3:
            question = match.group(1).strip()
            answer = match.group(2).strip()
            
            # Skip if already captured by previous patterns
            if any(qa_pairs[key]['question'] == question for key in qa_pairs):
                continue
            
            qa_count += 1
            qa_pairs[f"qa_{qa_count}"] = {
                'question': question,
                'answer': answer
            }
        
        return qa_pairs
    
    @staticmethod
    def _extract_qa_from_sections(text: str) -> Dict[str, Dict[str, Any]]:
        """
        Extract Q&A pairs from sections in academic documents.
        Returns a dictionary mapping question IDs to question-answer dictionaries.
        """
        # Identify potential section headers
        section_pattern = r'(?:^|\n)(?:\d+[\.)]\s*|\b(?:SECTION|PART|CHAPTER)\s*\d+:?\s+)([A-Z][A-Za-z\s]+)(?:\n|\:|$)'
        sections = re.finditer(section_pattern, text, re.MULTILINE)
        
        qa_pairs = {}
        qa_count = 0
        section_positions = []
        
        # First, find all section positions
        for match in sections:
            section_title = match.group(1).strip()
            if len(section_title) > 3 and len(section_title) < 100:  # Reasonable section title length
                section_positions.append((match.start(), match.end(), section_title))
        
        # No sections found
        if not section_positions:
            return {}
        
        # Get content between sections
        for i, (start, end, title) in enumerate(section_positions):
            if i < len(section_positions) - 1:
                next_start = section_positions[i + 1][0]
                section_content = text[end:next_start].strip()
            else:
                section_content = text[end:].strip()
            
            # Process content if it looks like it contains Q&A
            section_qa = DocxProcessor._extract_qa_from_section(title, section_content)
            if section_qa:
                for q, a in section_qa:
                    qa_count += 1
                    qa_pairs[f"qa_{qa_count}"] = {
                        "question": q,
                        "answer": a
                    }
        
        return qa_pairs
    
    @staticmethod
    def _extract_qa_from_section(section_title: str, content: str) -> List[Tuple[str, str]]:
        """
        Process a section's content to extract Q&A pairs.
        The section title may become a question if the content is an answer.
        """
        # Skip short or empty content
        if len(content) < 20:
            return []
        
        qa_pairs = []
        
        # Check if content contains subsections that look like questions
        question_pattern = r'(?:^|\n)(?:\d+[\.)]\s*|\([a-zA-Z0-9]+\)\s*)([^?]+(?:\?|\.))(?=\n)'
        questions = re.finditer(question_pattern, content, re.MULTILINE)
        
        questions_found = False
        for qmatch in questions:
            questions_found = True
            q_text = qmatch.group(1).strip()
            q_end = qmatch.end()
            
            # Find where this question's answer ends (at the next question or end of content)
            next_q_match = re.search(question_pattern, content[q_end:], re.MULTILINE)
            if next_q_match:
                a_text = content[q_end:q_end + next_q_match.start()].strip()
            else:
                a_text = content[q_end:].strip()
            
            qa_pairs.append((q_text, a_text))
        
        # If no subsection questions found, treat the section itself as a question
        if not questions_found and section_title and '?' not in section_title:
            # Check if the section title can be a question (e.g., "Introduction" probably isn't)
            if not re.match(r'\b(?:introduction|conclusion|summary|abstract|references|bibliography)\b', 
                           section_title, re.IGNORECASE):
                # Convert section title to a question if it's not already
                question = f"{section_title}?"
                qa_pairs.append((question, content))
        
        return qa_pairs 